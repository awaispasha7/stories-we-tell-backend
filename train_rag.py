#!/usr/bin/env python3
"""
Simple RAG Training Script

This single script handles everything:
- Auto-detects file types (DOCX, PDF, TXT)
- Converts files to text
- Trains the RAG system
- Shows progress and statistics

Usage:
    python train_rag.py "your-file.docx"
    python train_rag.py "your-file.pdf" 
    python train_rag.py "your-file.txt"
    python train_rag.py  # Interactive mode
"""

import asyncio
import argparse
import os
import sys
from pathlib import Path
from typing import List, Dict, Any, Optional

# Add the app directory to the path
sys.path.append(os.path.join(os.path.dirname(__file__), 'app'))

from dotenv import load_dotenv
load_dotenv()

# Try to import required libraries
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from app.ai.embedding_service import get_embedding_service
    from app.ai.vector_storage import vector_storage
    from app.database.supabase import get_supabase_client
    RAG_AVAILABLE = True
except ImportError as e:
    print(f"❌ RAG components not available: {e}")
    RAG_AVAILABLE = False


class SimpleRAGTrainer:
    """Simple RAG trainer that handles everything"""
    
    def __init__(self):
        if not RAG_AVAILABLE:
            raise Exception("RAG components not available. Check your setup.")
        
        self.embedding_service = get_embedding_service()
        self.vector_storage = vector_storage
        self.supabase = get_supabase_client()
    
    def detect_file_type(self, file_path: str) -> str:
        """Detect file type from extension"""
        ext = Path(file_path).suffix.lower()
        if ext == '.docx':
            return 'docx'
        elif ext == '.pdf':
            return 'pdf'
        elif ext in ['.txt', '.md']:
            return 'text'
        else:
            return 'unknown'
    
    def convert_to_text(self, file_path: str) -> str:
        """Convert any supported file to text"""
        file_type = self.detect_file_type(file_path)
        
        print(f"📄 Converting {file_type.upper()} file: {Path(file_path).name}")
        
        if file_type == 'docx':
            return self._convert_docx(file_path)
        elif file_type == 'pdf':
            return self._convert_pdf(file_path)
        elif file_type == 'text':
            return self._read_text(file_path)
        else:
            raise Exception(f"Unsupported file type: {file_type}")
    
    def _convert_docx(self, file_path: str) -> str:
        """Convert DOCX to text"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not installed. Run: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            raise Exception(f"Error converting DOCX: {e}")
    
    def _convert_pdf(self, file_path: str) -> str:
        """Convert PDF to text"""
        if not PDF_AVAILABLE:
            raise Exception("PyPDF2 not installed. Run: pip install PyPDF2")
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = []
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text.strip())
                
                return "\n\n".join(text_content)
                
        except Exception as e:
            raise Exception(f"Error converting PDF: {e}")
    
    def _read_text(self, file_path: str) -> str:
        """Read text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            raise Exception(f"Error reading text file: {e}")
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                search_start = max(start + chunk_size - 100, start)
                for i in range(end - 1, search_start, -1):
                    if text[i] in '.!?':
                        end = i + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            start = end - overlap
            if start >= len(text):
                break
        
        return chunks
    
    async def add_knowledge_chunk(self, content: str, source: str, knowledge_type: str = "global") -> bool:
        """Add a knowledge chunk to RAG"""
        try:
            # Generate embedding
            embedding = await self.embedding_service.generate_embedding(content)
            
            # Extract category from source filename
            source_name = Path(source).stem
            category = "general"
            if "story" in source_name.lower():
                category = "storytelling"
            elif "character" in source_name.lower():
                category = "character"
            elif "plot" in source_name.lower():
                category = "plot"
            elif "dialogue" in source_name.lower():
                category = "dialogue"
            
            # Store in vector database using existing function
            result = await self.vector_storage.store_global_knowledge(
                category=category,
                pattern_type="knowledge_chunk",
                embedding=embedding,
                example_text=content,
                description=f"Knowledge chunk from {source_name}",
                quality_score=0.8,
                tags=[source_name, "training"]
            )
            
            return result is not None
            
        except Exception as e:
            print(f"❌ Error adding chunk: {e}")
            return False
    
    async def train_from_file(self, file_path: str, knowledge_type: str = "global") -> bool:
        """Train RAG from any supported file"""
        try:
            # Convert file to text
            text_content = self.convert_to_text(file_path)
            
            print(f"📊 Content length: {len(text_content)} characters")
            
            # Split into chunks
            chunks = self.chunk_text(text_content)
            print(f"📦 Split into {len(chunks)} chunks")
            
            # Add each chunk to RAG
            success_count = 0
            for i, chunk in enumerate(chunks):
                print(f"🔄 Processing chunk {i+1}/{len(chunks)}")
                
                metadata = {
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "file_name": Path(file_path).name
                }
                
                success = await self.add_knowledge_chunk(
                    content=chunk,
                    source=file_path,
                    knowledge_type=knowledge_type
                )
                
                if success:
                    success_count += 1
                
                # Small delay to avoid rate limiting
                await asyncio.sleep(0.1)
            
            print(f"📈 Successfully processed {success_count}/{len(chunks)} chunks")
            return success_count > 0
            
        except Exception as e:
            print(f"❌ Error training from file: {e}")
            return False
    
    async def show_statistics(self):
        """Show knowledge base statistics"""
        try:
            result = self.supabase.table('global_knowledge').select('*').execute()
            
            if result.data:
                total_chunks = len(result.data)
                
                print(f"\n📊 Knowledge Base Statistics:")
                print(f"   Total knowledge chunks: {total_chunks}")
                
                # Show categories
                categories = {}
                for item in result.data:
                    category = item.get('category', 'Unknown')
                    categories[category] = categories.get(category, 0) + 1
                
                print(f"\n📁 Categories ({len(categories)}):")
                for category, count in sorted(categories.items()):
                    print(f"   - {category}: {count} chunks")
                
                # Show pattern types
                pattern_types = {}
                for item in result.data:
                    pattern_type = item.get('pattern_type', 'Unknown')
                    pattern_types[pattern_type] = pattern_types.get(pattern_type, 0) + 1
                
                print(f"\n🔍 Pattern Types ({len(pattern_types)}):")
                for pattern_type, count in sorted(pattern_types.items()):
                    print(f"   - {pattern_type}: {count} chunks")
            else:
                print("📊 No knowledge chunks found in database")
                
        except Exception as e:
            print(f"❌ Error retrieving statistics: {e}")
    
    async def interactive_mode(self):
        """Interactive mode for adding knowledge"""
        print("🎓 RAG Training - Interactive Mode")
        print("=" * 40)
        
        while True:
            print("\nOptions:")
            print("1. Train from file")
            print("2. Add text content")
            print("3. View statistics")
            print("4. Exit")
            
            choice = input("\nEnter your choice (1-4): ").strip()
            
            if choice == '1':
                file_path = input("Enter file path: ").strip()
                if os.path.exists(file_path):
                    knowledge_type = input("Knowledge type (global/user) [global]: ").strip() or "global"
                    success = await self.train_from_file(file_path, knowledge_type)
                    if success:
                        print("✅ File trained successfully!")
                    else:
                        print("❌ Failed to train file")
                else:
                    print("❌ File not found")
            
            elif choice == '2':
                source = input("Source/Title: ").strip() or "Interactive Input"
                print("Enter your content (press Ctrl+D or Ctrl+Z when done):")
                try:
                    content = ""
                    while True:
                        line = input()
                        content += line + "\n"
                except EOFError:
                    pass
                
                if content.strip():
                    knowledge_type = input("Knowledge type (global/user) [global]: ").strip() or "global"
                    success = await self.add_knowledge_chunk(content.strip(), source, knowledge_type)
                    if success:
                        print("✅ Content added successfully!")
                    else:
                        print("❌ Failed to add content")
                else:
                    print("❌ No content provided")
            
            elif choice == '3':
                await self.show_statistics()
            
            elif choice == '4':
                print("👋 Goodbye!")
                break
            
            else:
                print("❌ Invalid choice")


async def main():
    """Main function"""
    parser = argparse.ArgumentParser(description='Simple RAG Training Script')
    parser.add_argument('file', nargs='?', help='File to train with (DOCX, PDF, TXT)')
    parser.add_argument('--type', default='global', help='Knowledge type (global/user)')
    parser.add_argument('--interactive', action='store_true', help='Run in interactive mode')
    
    args = parser.parse_args()
    
    # Check environment
    required_vars = ['OPENAI_API_KEY', 'SUPABASE_URL', 'SUPABASE_SERVICE_ROLE_KEY']
    missing_vars = [var for var in required_vars if not os.getenv(var)]
    
    if missing_vars:
        print(f"❌ Missing environment variables: {', '.join(missing_vars)}")
        print("Please check your .env file")
        return
    
    if not RAG_AVAILABLE:
        print("❌ RAG components not available")
        print("Make sure you're in the backend directory and all dependencies are installed")
        return
    
    try:
        trainer = SimpleRAGTrainer()
        
        if args.interactive or not args.file:
            # Interactive mode
            await trainer.interactive_mode()
        else:
            # Train from file
            if not os.path.exists(args.file):
                print(f"❌ File not found: {args.file}")
                return
            
            print(f"🚀 Training RAG with: {Path(args.file).name}")
            success = await trainer.train_from_file(args.file, args.type)
            
            if success:
                print("✅ RAG training completed successfully!")
                await trainer.show_statistics()
            else:
                print("❌ RAG training failed")
    
    except Exception as e:
        print(f"❌ Error: {e}")
        print("\n💡 Make sure you have installed required dependencies:")
        print("   pip install python-docx PyPDF2")


if __name__ == "__main__":
    asyncio.run(main())
